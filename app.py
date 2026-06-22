import re
from io import BytesIO

import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Pt


st.set_page_config(page_title="AI Resume Optimizer", layout="centered")

st.title("🤖 AI Resume Optimizer & ATS Resume Builder")
st.write(
    "Upload a resume and paste a job description. The app analyzes skill match, "
    "optimizes the existing resume, and exports a polished DOCX."
)

skills = [
    "python", "java", "sql", "matlab", "aws", "git", "linux", "docker",
    "machine learning", "data analysis", "data visualization", "automation",
    "software development", "object-oriented programming", "oop", "testing",
    "debugging", "software maintenance", "application development",
    "requirements", "requirements analysis", "validation", "verification",
    "systems engineering", "industrial engineering", "simulation",
    "pandas", "numpy", "scikit-learn", "fastapi", "streamlit",
    "api", "rest api", "data pipelines", "documentation",
    "communication", "teamwork", "leadership", "problem solving",
    "qa", "quality assurance", "software testing", "troubleshooting",
    "javascript", "html", "css", "typescript", "spring boot",
    "postgresql", "ci/cd", "devops", "cloud"
]

for key, value in {
    "resume": "",
    "job": "",
    "target_role": "Software Engineer",
    "matched_skills": [],
    "missing_skills": [],
    "optimized_resume": "",
    "docx_file": None
}.items():
    if key not in st.session_state:
        st.session_state[key] = value


def clean_text(text):
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s+#./-]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text


def get_lines(text):
    return [line.strip() for line in text.splitlines() if line.strip()]


def find_skills(text, skills_list):
    text = clean_text(text)
    found = []

    for skill in skills_list:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, text):
            found.append(skill)

    return sorted(set(found))


def detect_role(job_text):
    job_lower = job_text.lower()

    role_keywords = [
        "software engineer",
        "software developer",
        "qa engineer",
        "quality assurance engineer",
        "test engineer",
        "ai engineer",
        "machine learning engineer",
        "ml engineer",
        "data engineer",
        "systems engineer",
        "data analyst",
        "data scientist",
        "cloud engineer",
        "devops engineer",
        "cybersecurity engineer",
        "backend engineer",
        "frontend engineer",
        "full stack engineer",
        "python developer",
        "java developer",
        "automation engineer",
        "simulation engineer"
    ]

    for role in role_keywords:
        if role in job_lower:
            return role.title()

    if "software" in job_lower:
        return "Software Engineer"
    if "qa" in job_lower or "testing" in job_lower:
        return "QA Engineer"
    if "data" in job_lower:
        return "Data Analyst"
    if "systems" in job_lower:
        return "Systems Engineer"
    if "machine learning" in job_lower or "ml" in job_lower:
        return "Machine Learning Engineer"
    if "ai" in job_lower:
        return "AI Engineer"

    return "Software Engineer"


def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(uploaded_file):
    text = ""
    doc = Document(uploaded_file)
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_txt(uploaded_file):
    return uploaded_file.read().decode("utf-8")


def extract_name(resume_text):
    lines = get_lines(resume_text)
    return lines[0] if lines else "Candidate Name"


def extract_contact(resume_text):
    lines = get_lines(resume_text)
    return lines[1] if len(lines) > 1 else "Location | Phone | Email | LinkedIn"


def section_between(resume_text, start_keywords, end_keywords):
    lines = get_lines(resume_text)
    start_index = None
    end_index = None

    for i, line in enumerate(lines):
        if line.upper() in start_keywords:
            start_index = i + 1
            break

    if start_index is None:
        return ""

    for j in range(start_index, len(lines)):
        if lines[j].upper() in end_keywords:
            end_index = j
            break

    if end_index is None:
        end_index = len(lines)

    return "\n".join(lines[start_index:end_index])


def extract_sections(resume_text):
    headers = [
        "SUMMARY",
        "PROFESSIONAL SUMMARY",
        "EDUCATION",
        "EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "PROJECTS",
        "TECHNICAL SKILLS",
        "SKILLS",
        "AWARDS"
    ]

    return {
        "summary": section_between(resume_text, ["SUMMARY", "PROFESSIONAL SUMMARY"], headers),
        "education": section_between(resume_text, ["EDUCATION"], headers),
        "experience": section_between(resume_text, ["EXPERIENCE", "PROFESSIONAL EXPERIENCE"], headers),
        "projects": section_between(resume_text, ["PROJECTS"], headers),
        "skills": section_between(resume_text, ["TECHNICAL SKILLS", "SKILLS"], headers),
        "awards": section_between(resume_text, ["AWARDS"], headers),
    }


def build_summary(target_role, resume_skills):
    if "Software" in target_role:
        return (
            "Software Engineer with experience developing Python-based applications, "
            "automating workflows, testing and debugging software solutions, and collaborating "
            "with engineering teams to deliver reliable technical solutions. Skilled in Python, "
            "Java, SQL, Linux, Git, Docker, object-oriented programming, and software development. "
            "Proven ability to improve operational efficiency through automation and deliver measurable business value."
        )

    if "AI" in target_role:
        return (
            "AI Engineer with experience building Python-based technical solutions, automating workflows, "
            "analyzing data, and applying machine learning tools to solve practical problems. Skilled in Python, "
            "machine learning, data analysis, automation, and software development."
        )

    if "Machine Learning" in target_role:
        return (
            "Machine Learning Engineer with experience using Python, Scikit-Learn, Pandas, NumPy, "
            "data analysis, automation, and model-based problem solving. Skilled in developing analytical workflows "
            "and applying machine learning techniques to real-world use cases."
        )

    if "Systems" in target_role:
        return (
            "Systems Engineer with experience supporting technical projects through requirements analysis, "
            "testing, validation, documentation, simulation, and process improvement. Skilled in systems thinking, "
            "problem solving, technical communication, and cross-functional collaboration."
        )

    return (
        f"{target_role} with experience developing technical solutions, automating workflows, "
        "testing and debugging applications, and collaborating with technical teams to deliver reliable results."
    )


def optimize_skills(skills_text):
    if not skills_text.strip():
        return (
            "Programming: Python, Java, SQL, MATLAB\n"
            "Software Development: Object-Oriented Programming (OOP), Software Testing, Debugging, Application Development, Software Maintenance\n"
            "Tools: Git, Linux, Docker, VS Code\n"
            "Data & Analytics: Pandas, NumPy, Scikit-Learn, Data Analysis, Automation"
        )

    text = skills_text

    text = text.replace("Programming: Java, Python, SQL, MATLAB", "Programming: Python, Java, SQL, MATLAB")
    text = text.replace("objectoriented", "object-oriented")
    text = text.replace("ObjectOriented", "Object-Oriented")
    text = text.replace("Application\nDevelopment", "Application Development")

    if "Automation" not in text and "automation" not in text:
        text += "\nAutomation: Workflow Automation, Data Validation, Process Improvement"

    return text


def optimize_bullet(bullet):
    clean = bullet.replace("•", "").replace("-", "").strip()
    lower = clean.lower()

    if not clean:
        return ""

    clean = clean.replace("objectoriented", "object-oriented")
    clean = clean.replace("dataprocessing", "data-processing")
    clean = clean.replace("errorhandling", "error-handling")
    clean = clean.replace("realtime", "real-time")
    clean = clean.replace("simulationbased", "simulation-based")

    if "reduced processing time by 40%" in lower:
        return "Built automation tools that reduced processing time by 40%, improving workflow efficiency and reducing manual effort."

    if "tested" in lower or "debugged" in lower:
        return "Tested, debugged, and maintained software solutions to improve reliability, usability, and performance."

    if "requirements" in lower:
        return "Collaborated with engineers to translate requirements into practical technical solutions."

    if "developed software applications" in lower:
        return "Developed software applications supporting operational workflows and business processes."

    if "developed software tools" in lower:
        return "Developed Python-based software tools using object-oriented programming principles."

    if "automated reporting" in lower:
        return "Automated reporting and data-processing workflows, improving consistency and reducing manual work."

    if "developed automation solutions" in lower:
        return "Developed automation solutions for processing and validating large datasets."

    if "validation" in lower or "error" in lower:
        return "Implemented validation checks and error-handling functionality to improve data quality and reliability."

    if clean.endswith("."):
        return clean

    return clean + "."


def parse_experience(experience_text):
    lines = get_lines(experience_text)
    jobs = []
    current_job = None

    for line in lines:
        if line.startswith("•") or line.startswith("-"):
            if current_job:
                bullet = optimize_bullet(line)
                if bullet:
                    current_job["bullets"].append(bullet)
        else:
            current_job = {"title": line, "bullets": []}
            jobs.append(current_job)

    return jobs


def parse_projects(projects_text):
    lines = get_lines(projects_text)
    projects = []
    current_project = None

    for line in lines:
        if line.startswith("•") or line.startswith("-"):
            if current_project:
                bullet = optimize_bullet(line)
                if bullet:
                    current_project["bullets"].append(bullet)
        else:
            current_project = {"name": line, "bullets": []}
            projects.append(current_project)

    return projects


def improve_project_bullets(project):
    name = project["name"]
    bullets = project["bullets"]

    if "UGV" in name:
        return [
            "Developed a real-time localization and control platform using Python.",
            "Improved navigation accuracy through algorithm optimization, testing, and performance refinement.",
            "Evaluated software performance through iterative development and system validation."
        ]

    if "Automation Utilities" in name:
        return [
            "Developed automation tools for data validation, reporting, and workflow improvement.",
            "Reduced manual processing requirements by improving repeatability and data reliability."
        ]

    if "Humanitarian Aerial Logistics" in name:
        return [
            "Developed simulation-based logistics models to evaluate operational feasibility and system performance.",
            "Performed data analysis and system evaluation to support operational decision-making.",
            "Presented technical findings and recommendations to stakeholders."
        ]

    return bullets


def generate_optimized_resume(resume_text, job_text):
    sections = extract_sections(resume_text)

    name = extract_name(resume_text)
    contact = extract_contact(resume_text)

    target_role = detect_role(job_text)
    resume_skills = find_skills(resume_text, skills)
    job_skills = find_skills(job_text, skills)

    matched_skills = sorted(set(resume_skills).intersection(set(job_skills)))
    missing_skills = sorted(set(job_skills).difference(set(resume_skills)))

    st.session_state.target_role = target_role
    st.session_state.matched_skills = matched_skills
    st.session_state.missing_skills = missing_skills

    summary = build_summary(target_role, resume_skills)
    skills_section = optimize_skills(sections["skills"])

    experience = parse_experience(sections["experience"])
    projects = parse_projects(sections["projects"])

    for project in projects:
        project["bullets"] = improve_project_bullets(project)

    return {
        "name": name,
        "contact": contact,
        "target_role": target_role,
        "summary": summary,
        "skills": skills_section,
        "experience": experience,
        "projects": projects,
        "education": sections["education"],
        "awards": sections["awards"],
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "resume_skills": resume_skills,
        "job_skills": job_skills
    }


def resume_to_markdown(data):
    lines = []

    lines.append(f"# {data['name']}")
    lines.append(data["contact"])

    lines.append("\n## PROFESSIONAL SUMMARY")
    lines.append(data["summary"])

    lines.append("\n## TECHNICAL SKILLS")
    lines.append(data["skills"])

    lines.append("\n## PROFESSIONAL EXPERIENCE")
    for job in data["experience"]:
        lines.append(f"\n**{job['title']}**")
        for bullet in job["bullets"]:
            lines.append(f"- {bullet}")

    lines.append("\n## PROJECTS")
    for project in data["projects"]:
        lines.append(f"\n**{project['name']}**")
        for bullet in project["bullets"]:
            lines.append(f"- {bullet}")

    if data["education"].strip():
        lines.append("\n## EDUCATION")
        lines.append(data["education"])

    if data["awards"].strip():
        lines.append("\n## AWARDS")
        lines.append(data["awards"])

    return "\n".join(lines)


def create_docx(data):
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading(data["name"].upper(), level=0)
    doc.add_paragraph(data["contact"])

    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(data["summary"])

    doc.add_heading("TECHNICAL SKILLS", level=1)
    for line in get_lines(data["skills"]):
        doc.add_paragraph(line)

    doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    for job in data["experience"]:
        doc.add_paragraph(job["title"])
        for bullet in job["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("PROJECTS", level=1)
    for project in data["projects"]:
        doc.add_paragraph(project["name"])
        for bullet in project["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    if data["education"].strip():
        doc.add_heading("EDUCATION", level=1)
        for line in get_lines(data["education"]):
            doc.add_paragraph(line)

    if data["awards"].strip():
        doc.add_heading("AWARDS", level=1)
        for line in get_lines(data["awards"]):
            doc.add_paragraph(line, style="List Bullet")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


uploaded_resume = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])

if uploaded_resume:
    if uploaded_resume.name.endswith(".pdf"):
        st.session_state.resume = extract_text_from_pdf(uploaded_resume)
    elif uploaded_resume.name.endswith(".docx"):
        st.session_state.resume = extract_text_from_docx(uploaded_resume)
    elif uploaded_resume.name.endswith(".txt"):
        st.session_state.resume = extract_text_from_txt(uploaded_resume)

    st.success("Resume uploaded successfully.")

    with st.expander("View Extracted Resume Text"):
        st.text(st.session_state.resume[:3000])


job = st.text_area("Job Description", height=250)
st.session_state.job = job

if job.strip():
    st.info(f"Detected Role: {detect_role(job)}")


if st.button("Analyze Resume"):
    if not st.session_state.resume.strip() or not st.session_state.job.strip():
        st.warning("Please upload a resume and paste a job description.")
    else:
        resume_skills = find_skills(st.session_state.resume, skills)
        job_skills = find_skills(st.session_state.job, skills)

        matched_skills = sorted(set(resume_skills).intersection(set(job_skills)))
        missing_skills = sorted(set(job_skills).difference(set(resume_skills)))

        score = round((len(matched_skills) / len(job_skills)) * 100, 2) if job_skills else 0

        st.subheader("Resume Analysis Results")
        st.metric("Skill Match Score", f"{score}%")

        col1, col2, col3 = st.columns(3)
        col1.metric("Resume Skills Found", len(resume_skills))
        col2.metric("Job Skills Required", len(job_skills))
        col3.metric("Matched Skills", len(matched_skills))

        st.write("### ✅ Matched Skills")
        st.success(", ".join(matched_skills) if matched_skills else "No matched skills found.")

        st.write("### ⚠️ Missing Skills")
        st.warning(", ".join(missing_skills) if missing_skills else "No major missing skills found.")


st.divider()

st.write("## 📄 Generate Optimized Resume")
st.write("Optimizes your existing resume and exports a polished DOCX.")

if st.button("Generate Optimized Resume"):
    if not st.session_state.resume.strip() or not st.session_state.job.strip():
        st.warning("Please upload a resume and paste a job description first.")
    else:
        data = generate_optimized_resume(
            st.session_state.resume,
            st.session_state.job
        )

        st.session_state.optimized_resume = resume_to_markdown(data)
        st.session_state.docx_file = create_docx(data)

if st.session_state.optimized_resume:
    st.markdown(st.session_state.optimized_resume)

if st.session_state.docx_file:
    file_name = f"{st.session_state.target_role.replace(' ', '_')}_Resume.docx"

    st.download_button(
        label="📄 Download Optimized Resume as DOCX",
        data=st.session_state.docx_file,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
